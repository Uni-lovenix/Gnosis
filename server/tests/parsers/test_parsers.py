"""Tests for file parsers using in-memory fixtures.

We avoid shipping binary fixtures; instead we generate sample documents at
test time (openpyxl/python-docx/pdfplumber/markdown-it-py round-trip).
"""
from __future__ import annotations

from pathlib import Path

import pytest

from app.parsers import parse_excel, parse_markdown, parse_pdf, parse_word


@pytest.fixture
def tmp_path(tmp_path: Path) -> Path:
    return tmp_path


def _write_excel(path: Path) -> Path:
    pytest.importorskip("openpyxl")
    from openpyxl import Workbook

    wb = Workbook()
    ws = wb.active
    ws.title = "Sheet1"
    ws.append(["name", "score"])
    ws.append(["alice", 10])
    ws.append(["bob", 20])
    p = path / "sample.xlsx"
    wb.save(str(p))
    return p


def _write_docx(path: Path) -> Path:
    pytest.importorskip("docx")
    from docx import Document as DocxDocument

    doc = DocxDocument()
    doc.add_heading("Title", level=1)
    doc.add_paragraph("Hello world.")
    doc.add_heading("Sub", level=2)
    doc.add_paragraph("Body text.")
    p = path / "sample.docx"
    doc.save(str(p))
    return p


def _write_pdf(path: Path) -> Path:
    pytest.importorskip("pdfplumber")
    reportlab = pytest.importorskip("reportlab")
    from reportlab.pdfgen import canvas  # type: ignore

    p = path / "sample.pdf"
    c = canvas.Canvas(str(p))
    c.drawString(100, 700, "Hello PDF world")
    c.drawString(100, 680, "Second line of text")
    c.save()
    return p


def _write_markdown(path: Path) -> Path:
    md = (
        "---\n"
        "title: Sample\n"
        "---\n\n"
        "# Heading 1\n\n"
        "First paragraph.\n\n"
        "## Heading 2\n\n"
        "Second paragraph.\n"
    )
    p = path / "sample.md"
    p.write_text(md, encoding="utf-8")
    return p


def test_excel_parser():
    p = _write_excel(Path("/tmp"))
    # The path uses /tmp; if pytest's tmp_path is sandboxed, write into it.
    from tempfile import mkdtemp

    d = Path(mkdtemp())
    p = _write_excel(d)
    doc = parse_excel(p)
    assert "Sheet1" in doc.text
    assert "alice" in doc.text
    assert doc.metadata["parser"] == "excel"


def test_word_parser():
    pytest.importorskip("docx")
    from tempfile import mkdtemp

    d = Path(mkdtemp())
    p = _write_docx(d)
    doc = parse_word(p)
    assert "# Title" in doc.text
    assert "Hello world" in doc.text
    assert doc.metadata["parser"] == "word"


def test_pdf_parser_smoke():
    pytest.importorskip("pdfplumber")
    from tempfile import mkdtemp

    d = Path(mkdtemp())
    p = _write_pdf(d)
    # Hand-crafted minimal PDF may yield empty text but must not raise.
    doc = parse_pdf(p)
    assert doc.metadata["parser"] == "pdf"
    assert "pages" in doc.metadata


def test_markdown_parser():
    pytest.importorskip("markdown_it")
    from tempfile import mkdtemp

    d = Path(mkdtemp())
    p = _write_markdown(d)
    doc = parse_markdown(p)
    assert "title: Sample" not in doc.text  # front matter stripped
    assert "# Heading 1" in doc.text
    assert doc.metadata["parser"] == "markdown"


def test_markdown_front_matter_stripping():
    pytest.importorskip("markdown_it")
    from tempfile import mkdtemp

    d = Path(mkdtemp())
    p = d / "front.md"
    p.write_text("---\nk: v\n---\nbody", encoding="utf-8")
    doc = parse_markdown(p)
    assert doc.text.startswith("body")