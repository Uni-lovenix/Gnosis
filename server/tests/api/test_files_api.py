"""API tests for the files router (multipart upload + task status).

Imports are dispatched via FastAPI ``BackgroundTasks`` (Fix-KB-Upload-
Progress); the POST returns ``task_id`` immediately and the pipeline runs
in the background. Tests therefore assert on the eventual task state by
polling the GET endpoint rather than the POST response body.
"""
from __future__ import annotations

import time
from pathlib import Path

import pytest
from fastapi.testclient import TestClient

from app.api import files as files_api
from app.main import create_app
from app.observability.task_store import TaskStore


@pytest.fixture
def client(tmp_path: Path):
    app = create_app()
    files_api.set_task_store(TaskStore(tmp_path / "tasks.db"))
    return TestClient(app)


@pytest.fixture
def excel_bytes():
    pytest.importorskip("openpyxl")
    from openpyxl import Workbook
    from io import BytesIO

    wb = Workbook()
    ws = wb.active
    ws.append(["a", "b"])
    ws.append([1, 2])
    buf = BytesIO()
    wb.save(buf)
    return buf.getvalue()


@pytest.fixture
def markdown_bytes():
    return b"# Title\n\nbody"


def _wait_for_task(client, task_id: str, timeout_s: float = 5.0) -> dict:
    """Poll the task endpoint until status reaches a terminal state.

    Tests run in-process via TestClient, so the BackgroundTasks runner drains
    almost immediately, but we still allow up to ``timeout_s`` for the result
    to settle (in case of slow embedder init).
    """
    deadline = time.monotonic() + timeout_s
    while time.monotonic() < deadline:
        r = client.get(f"/v1/files/tasks/{task_id}")
        assert r.status_code == 200
        body = r.json()
        if body["status"] in {"done", "failed"}:
            return body
        time.sleep(0.05)
    raise AssertionError(f"task {task_id} did not terminate within {timeout_s}s")


def test_import_markdown(client, markdown_bytes):
    r = client.post(
        "/v1/files/import",
        files={"file": ("a.md", markdown_bytes, "text/markdown")},
    )
    assert r.status_code == 200
    body = r.json()
    # POST returns placeholder counters (FillKB-Async-Import); real values
    # are written by the background pipeline.
    assert body["chunks"] == 0
    assert body["parser"] is None
    assert body["task_id"]

    final = _wait_for_task(client, body["task_id"])
    assert final["status"] == "done"
    assert final["progress"] == 1.0
    assert final["result"]["chunks"] >= 1
    assert final["result"]["parser"] == "markdown"


def test_import_excel(client, excel_bytes):
    r = client.post(
        "/v1/files/import",
        files={"file": ("a.xlsx", excel_bytes, "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")},
    )
    assert r.status_code == 200
    body = r.json()
    assert body["task_id"]
    final = _wait_for_task(client, body["task_id"])
    assert final["status"] == "done"
    assert final["result"]["parser"] == "excel"
    assert final["result"]["chunks"] >= 1


def test_unsupported_type(client):
    r = client.post(
        "/v1/files/import",
        files={"file": ("a.xyz", b"hello", "application/octet-stream")},
    )
    assert r.status_code == 415


def test_empty_file(client):
    r = client.post(
        "/v1/files/import",
        files={"file": ("a.md", b"", "text/markdown")},
    )
    assert r.status_code == 400


def test_task_not_found(client):
    r = client.get("/v1/files/tasks/missing")
    assert r.status_code == 404


def test_word_import(client):
    pytest.importorskip("docx")
    from docx import Document as DocxDocument
    from io import BytesIO

    doc = DocxDocument()
    doc.add_heading("T", level=1)
    doc.add_paragraph("body")
    buf = BytesIO()
    doc.save(buf)
    r = client.post(
        "/v1/files/import",
        files={"file": ("a.docx", buf.getvalue(), "application/vnd.openxmlformats-officedocument.wordprocessingml.document")},
    )
    assert r.status_code == 200
    body = r.json()
    final = _wait_for_task(client, body["task_id"])
    assert final["status"] == "done"
    assert final["result"]["parser"] == "word"
