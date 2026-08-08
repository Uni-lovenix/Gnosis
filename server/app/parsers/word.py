"""Word parser: python-docx → Document.

Each paragraph becomes a line; headings are prefixed with `#` according to
their outline level. Tables are rendered as pipe tables.
"""
from __future__ import annotations

from pathlib import Path

from app.observability.models import Document
from app.parsers._common import detect_mime


def parse_word(path: str | Path) -> Document:
    try:
        from docx import Document as DocxDocument  # type: ignore
    except ImportError as e:  # pragma: no cover
        raise RuntimeError(
            "python-docx not installed. `pip install -e '.[parsers]'`."
        ) from e

    p = Path(path)
    doc = DocxDocument(str(p))
    parts: list[str] = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            parts.append("")
            continue
        style = (para.style.name or "").lower()
        if "heading" in style:
            level = 1
            for n in range(1, 7):
                if f"heading {n}" in style:
                    level = n
                    break
            parts.append("#" * level + " " + text)
        else:
            parts.append(text)

    for table in doc.tables:
        rows = [[cell.text.strip() for cell in row.cells] for row in table.rows]
        if not rows:
            continue
        parts.append("| " + " | ".join(rows[0]) + " |")
        parts.append("|" + "|".join(["---"] * len(rows[0])) + "|")
        for r in rows[1:]:
            parts.append("| " + " | ".join(r) + " |")
        parts.append("")

    return Document(
        source_path=str(p),
        mime=detect_mime(p),
        text="\n".join(parts).strip(),
        metadata={
            "parser": "word",
            "paragraphs": len(doc.paragraphs),
            "tables": len(doc.tables),
            "size": p.stat().st_size,
        },
    )